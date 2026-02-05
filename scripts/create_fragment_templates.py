"""
Create test fragment templates for modular pleading generation.

Run this once to create the fragment .docx files.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FRAGMENTS_DIR = Path(__file__).parent.parent / "templates" / "pleadings" / "fragments"
FRAGMENTS_DIR.mkdir(exist_ok=True)


def create_notice_fragment():
    """Create Notice of Motion fragment."""
    doc = Document()

    # Title
    title = doc.add_paragraph("NOTICE OF MOTION")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(12)

    # Body
    doc.add_paragraph()
    doc.add_paragraph("TO ALL PARTIES AND THEIR ATTORNEYS OF RECORD:")
    doc.add_paragraph()
    p = doc.add_paragraph("PLEASE TAKE NOTICE that on ")
    p.add_run("{{ hearing_date }}").bold = True
    p.add_run(" at ")
    p.add_run("{{ hearing_time }}").bold = True
    p.add_run(", or as soon thereafter as the matter may be heard, in ")
    p.add_run("{{ hearing_location }}").bold = True
    p.add_run(" of the above-entitled Court, {{ document_title }} will be heard.")

    doc.add_paragraph()
    doc.add_paragraph("This motion is made pursuant to [applicable rules] and is based on this Notice of Motion, the attached Memorandum of Points and Authorities, the Declaration(s) filed herewith, and such other matters as may be presented at the hearing.")

    doc.save(FRAGMENTS_DIR / "notice.docx")
    print("Created: notice.docx")


def create_meet_confer_fragment():
    """Create Meet and Confer fragment."""
    doc = Document()

    # Title
    title = doc.add_paragraph("MEET AND CONFER DECLARATION")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(12)

    # Body
    doc.add_paragraph()
    doc.add_paragraph("Pursuant to Local Rule 7-3, counsel for {{ plaintiff_caption }} declares as follows:")
    doc.add_paragraph()
    doc.add_paragraph("1. Prior to filing this motion, counsel attempted in good faith to meet and confer with opposing counsel regarding the issues raised herein.")
    doc.add_paragraph()
    doc.add_paragraph("2. [Details of meet and confer efforts to be added]")

    doc.save(FRAGMENTS_DIR / "meet_confer.docx")
    print("Created: meet_confer.docx")


def create_toc_fragment():
    """Create Table of Contents fragment."""
    doc = Document()

    # Title
    title = doc.add_paragraph("TABLE OF CONTENTS")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(12)

    # Placeholder content
    doc.add_paragraph()
    doc.add_paragraph("I. INTRODUCTION....................................................1")
    doc.add_paragraph("II. STATEMENT OF FACTS.............................................2")
    doc.add_paragraph("III. LEGAL STANDARD................................................3")
    doc.add_paragraph("IV. ARGUMENT.......................................................4")
    doc.add_paragraph("V. CONCLUSION......................................................5")
    doc.add_paragraph()

    doc.save(FRAGMENTS_DIR / "toc.docx")
    print("Created: toc.docx")


def create_toa_fragment():
    """Create Table of Authorities fragment."""
    doc = Document()

    # Title
    title = doc.add_paragraph("TABLE OF AUTHORITIES")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(12)

    # Cases header
    doc.add_paragraph()
    cases = doc.add_paragraph("CASES")
    cases.runs[0].bold = True
    cases.runs[0].underline = True

    doc.add_paragraph()
    doc.add_paragraph("[Case citations to be added].........................................1")
    doc.add_paragraph()

    # Statutes header
    statutes = doc.add_paragraph("STATUTES")
    statutes.runs[0].bold = True
    statutes.runs[0].underline = True

    doc.add_paragraph()
    doc.add_paragraph("[Statutory citations to be added]....................................2")
    doc.add_paragraph()

    doc.save(FRAGMENTS_DIR / "toa.docx")
    print("Created: toa.docx")


def create_cert_compliance_fragment():
    """Create Certificate of Compliance fragment."""
    doc = Document()

    # Title
    title = doc.add_paragraph("CERTIFICATE OF COMPLIANCE")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(12)

    # Body
    doc.add_paragraph()
    doc.add_paragraph("Pursuant to Local Rule 7-4, the undersigned certifies that this memorandum complies with the applicable word limit.")
    doc.add_paragraph()
    doc.add_paragraph("This memorandum contains approximately [WORD COUNT] words, excluding the portions exempted by Local Rule 7-4.")
    doc.add_paragraph()

    # Signature
    doc.add_paragraph()
    doc.add_paragraph("Dated: {{ date_formatted }}")
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph("_______________________________")
    doc.add_paragraph("{{ associate_name }}")
    doc.add_paragraph("Attorney for Plaintiffs")

    doc.save(FRAGMENTS_DIR / "cert_compliance.docx")
    print("Created: cert_compliance.docx")


def create_base_with_subdocs():
    """Create a base template that uses subdoc placeholders."""
    doc = Document()

    # Caption area
    title = doc.add_paragraph("{{ court_name }}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Case No.: {{ case_number }}")
    doc.add_paragraph()
    doc.add_paragraph("{{ plaintiff_caption }},")
    doc.add_paragraph("        Plaintiffs,")
    doc.add_paragraph("    v.")
    doc.add_paragraph("{{ defendant_caption }},")
    doc.add_paragraph("        Defendants.")
    doc.add_paragraph()

    # Document title
    doc_title = doc.add_paragraph("{{ document_title }}")
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_title.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph("Judge: {{ judge_name }}")
    doc.add_paragraph()

    # Subdoc placeholders - these get replaced with actual subdocs
    # The {% if %} conditionals control whether they're included
    doc.add_paragraph("{% if include_notice %}{{ subdoc_notice }}{% endif %}")
    doc.add_paragraph("{% if include_meet_confer %}{{ subdoc_meet_confer }}{% endif %}")
    doc.add_paragraph("{% if include_toc %}{{ subdoc_toc }}{% endif %}")
    doc.add_paragraph("{% if include_toa %}{{ subdoc_toa }}{% endif %}")

    # Memo body placeholder
    doc.add_paragraph()
    memo_title = doc.add_paragraph("MEMORANDUM OF POINTS AND AUTHORITIES")
    memo_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    memo_title.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph("I. INTRODUCTION")
    doc.add_paragraph("{{ introduction }}")
    doc.add_paragraph()
    doc.add_paragraph("II. LEGAL STANDARD")
    doc.add_paragraph("{{ legal_standard }}")
    doc.add_paragraph()
    doc.add_paragraph("III. ARGUMENT")
    doc.add_paragraph("{{ argument }}")
    doc.add_paragraph()
    doc.add_paragraph("IV. CONCLUSION")
    doc.add_paragraph("{{ conclusion }}")
    doc.add_paragraph()

    # Certificate of compliance subdoc
    doc.add_paragraph("{% if include_cert_compliance %}{{ subdoc_cert_compliance }}{% endif %}")

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("Dated: {{ date_formatted }}")
    doc.add_paragraph()
    doc.add_paragraph("{{ firm_name }}")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("By: _______________________________")
    doc.add_paragraph("    {{ associate_name }}")
    doc.add_paragraph("    State Bar No. {{ bar_number }}")
    doc.add_paragraph("    {{ attorney_email }}")
    doc.add_paragraph("    Attorney for Plaintiffs")

    # Save
    base_path = Path(__file__).parent.parent / "templates" / "pleadings" / "base_subdoc.docx"
    doc.save(base_path)
    print(f"Created: base_subdoc.docx")


if __name__ == "__main__":
    print("Creating fragment templates...")
    create_notice_fragment()
    create_meet_confer_fragment()
    create_toc_fragment()
    create_toa_fragment()
    create_cert_compliance_fragment()
    create_base_with_subdocs()
    print("\nDone! Fragment templates created in templates/pleadings/fragments/")
