"""
Build the DOCX template for case report export.

Generates templates/exports/case_report_template.docx programmatically
using python-docx. The template contains Jinja2 tags that docxtpl renders
at export time with real case data and RichText objects.

Run:  python scripts/build_docx_template.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUTPUT = Path(__file__).parent.parent / "templates" / "exports" / "case_report_template.docx"

# Colors — refined palette
NAVY = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x94, 0xA3, 0xB8)
DARK_GREY = RGBColor(0x47, 0x55, 0x69)
BLUE = RGBColor(0x3B, 0x82, 0xF6)       # blue-500 accent
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
WARM_WHITE = RGBColor(0xFA, 0xFA, 0xF9)  # stone-50

NAVY_HEX = "0F172A"
HEADER_BG = "1E293B"    # slate-800 — softer than pure navy for table headers
ACCENT_HEX = "3B82F6"   # blue-500 — accent rules and highlights
GREY_BORDER = "E2E8F0"
FINE_BORDER = "CBD5E1"   # slate-300 — for subtle structural lines
LIGHT_BG = "F8FAFC"      # slate-50


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each arg is (color_hex, size_eighths)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            color, sz = val
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), color)
            el.set(qn("w:space"), "0")
            borders.append(el)
    tc_pr.append(borders)


def set_cell_margins(cell, top=0, bottom=0, left=60, right=60):
    """Set cell margins in twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)


def set_cell_width(cell, inches):
    """Set preferred cell width."""
    tc_pr = cell._tc.get_or_add_tcPr()
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(int(inches * 1440)))
    width.set(qn("w:type"), "dxa")
    tc_pr.append(width)


def set_table_borders(table, color="E2E8F0", sz=4):
    """Set uniform borders on entire table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def add_run(para, text, size=None, bold=False, italic=False, color=None, font_name=None):
    """Add a styled run to a paragraph."""
    run = para.add_run(text)
    if size:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    else:
        run.font.name = "Calibri"
    return run


def set_paragraph_spacing(para, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line:
        fmt.line_spacing = Pt(line)


def set_paragraph_border_bottom(para, color="E2E8F0", sz=4):
    """Add a bottom border to a paragraph."""
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:space"), "1")
    borders.append(bottom)
    p_pr.append(borders)


def set_page_break_before(para):
    """Set page break before paragraph."""
    p_pr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    p_pr.append(pb)


def build_header_footer(doc):
    """Add header and footer to the default section."""
    section = doc.sections[0]
    # Header: right-aligned, 7.5pt dark grey
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(hp, "{{ attorney_name }}", size=7.5, bold=True, color=DARK_GREY)
    add_run(hp, "  |  ", size=7.5, color=GREY)
    add_run(hp, "{{ creation_date }}", size=7.5, color=GREY)

    # Footer: right-aligned page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run()
    run.font.size = Pt(7)
    run.font.color.rgb = GREY
    run.font.name = "Calibri"
    # Insert PAGE field
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def build_cover_page(doc):
    """Build the cover page with title and phase-grouped summary table."""
    # Spacer for top margin breathing room
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=12)

    # Title — large, commanding
    title_p = doc.add_paragraph()
    set_paragraph_spacing(title_p, after=3)
    add_run(title_p, "Active Case Portfolio", size=26, bold=True, color=NAVY)

    # Blue accent rule
    rule_p = doc.add_paragraph()
    set_paragraph_spacing(rule_p, after=4)
    set_paragraph_border_bottom(rule_p, color=ACCENT_HEX, sz=10)

    # Subtitle — attorney + date
    sub_p = doc.add_paragraph()
    set_paragraph_spacing(sub_p, after=1)
    add_run(sub_p, "{{ attorney_name }}", size=11, color=NAVY)
    add_run(sub_p, "  |  ", size=11, color=GREY)
    add_run(sub_p, "{{ creation_date }}", size=11, color=DARK_GREY)

    # Case count
    count_p = doc.add_paragraph()
    set_paragraph_spacing(count_p, after=10)
    add_run(count_p, "{{ case_count }} active matters", size=9, color=GREY)

    # Summary table with phase grouping
    table = doc.add_table(rows=7, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color=FINE_BORDER, sz=2)

    col_widths = [2.0, 1.2, 1.1, 1.4, 0.9]

    # Row 0: Column headers — softer slate background
    headers = ["Case", "Case No.", "Judge", "Clients", "DOI"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell, top=35, bottom=35, left=55, right=55)
        set_cell_width(cell, col_widths[i])
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        add_run(p, h, size=7, bold=True, color=WHITE)

    # Row 1: {%tr for phase in phases %}
    p = table.cell(1, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr for phase in phases %}", size=8)
    for i in range(1, 5):
        set_paragraph_spacing(table.cell(1, i).paragraphs[0], before=0, after=0)

    # Row 2: Phase header (merge all cells, colored bg)
    merged_cell = table.cell(2, 0).merge(table.cell(2, 4))
    set_cell_margins(merged_cell, top=30, bottom=30, left=55, right=55)
    p = merged_cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{% cellbg phase.bg %}", size=1, color=WHITE)
    add_run(p, "{{r phase.header_rt }}", size=10)

    # Row 3: {%tr for case in phase.cases %}
    p = table.cell(3, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr for case in phase.cases %}", size=8)
    for i in range(1, 5):
        set_paragraph_spacing(table.cell(3, i).paragraphs[0], before=0, after=0)

    # Row 4: Data row — more padding for readability
    data_tags = [
        ("{{ case.short_name }}", 8, True, None, "Calibri"),
        ("{{ case.case_number }}", 7.5, False, None, "Consolas"),
        ("{{ case.judge }}", 8, False, None, "Calibri"),
        ("{{ case.clients }}", 8, False, None, "Calibri"),
        ("{{ case.doi }}", 8, False, None, "Calibri"),
    ]
    for i, (tag, sz, bold, clr, font) in enumerate(data_tags):
        cell = table.cell(4, i)
        set_cell_margins(cell, top=25, bottom=25, left=55, right=55)
        set_cell_width(cell, col_widths[i])
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        add_run(p, tag, size=sz, bold=bold, font_name=font)

    # Row 5: {%tr endfor %} (inner: cases)
    p = table.cell(5, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr endfor %}", size=8)
    for i in range(1, 5):
        set_paragraph_spacing(table.cell(5, i).paragraphs[0], before=0, after=0)

    # Row 6: {%tr endfor %} (outer: phases)
    p = table.cell(6, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr endfor %}", size=8)
    for i in range(1, 5):
        set_paragraph_spacing(table.cell(6, i).paragraphs[0], before=0, after=0)


def build_section_heading(doc, text):
    """Add a section heading like 'KEY DATES & EVENTS'."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    set_paragraph_border_bottom(p, color=ACCENT_HEX, sz=4)
    add_run(p, text, size=9, bold=True, color=NAVY)


def build_events_table(doc):
    """Build the events table with {%tr for} loop and RichText + cellbg."""
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color=FINE_BORDER, sz=2)

    col_widths = [1.5, 3.5, 1.6]
    headers = ["Date", "Description", "Location"]

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell, top=30, bottom=30, left=55, right=55)
        set_cell_width(cell, col_widths[i])
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        add_run(p, h, size=7, bold=True, color=WHITE)

    # {%tr for event in case.events %}
    p = table.cell(1, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr for event in case.events %}", size=8)
    for i in range(1, 3):
        set_paragraph_spacing(table.cell(1, i).paragraphs[0], before=0, after=0)

    # Data row with RichText tags + cellbg
    data_cells = [
        ("{{r event.date_rt }}", 1.5),
        ("{{r event.desc_rt }}", 3.5),
        ("{{r event.loc_rt }}", 1.6),
    ]
    for i, (tag, w) in enumerate(data_cells):
        cell = table.cell(2, i)
        set_cell_margins(cell, top=15, bottom=15, left=50, right=50)
        set_cell_width(cell, w)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        # Add cellbg tag first in the cell
        add_run(p, "{% cellbg event.bg %}", size=8)
        add_run(p, tag, size=8)

    # {%tr endfor %}
    p = table.cell(3, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr endfor %}", size=8)
    for i in range(1, 3):
        set_paragraph_spacing(table.cell(3, i).paragraphs[0], before=0, after=0)


def build_tasks_table(doc):
    """Build the tasks table with {%tr for} loop and RichText + cellbg."""
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color=FINE_BORDER, sz=2)

    col_widths = [1.5, 3.8, 1.3]
    headers = ["Due Date", "Description", "Urgency"]

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell, top=30, bottom=30, left=55, right=55)
        set_cell_width(cell, col_widths[i])
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        add_run(p, h, size=7, bold=True, color=WHITE)

    # {%tr for task in case.tasks %}
    p = table.cell(1, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr for task in case.tasks %}", size=8)
    for i in range(1, 3):
        set_paragraph_spacing(table.cell(1, i).paragraphs[0], before=0, after=0)

    # Data row
    data_cells = [
        ("{{r task.date_rt }}", 1.5),
        ("{{r task.desc_rt }}", 3.8),
        ("{{r task.urgency_rt }}", 1.3),
    ]
    for i, (tag, w) in enumerate(data_cells):
        cell = table.cell(2, i)
        set_cell_margins(cell, top=15, bottom=15, left=50, right=50)
        set_cell_width(cell, w)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        add_run(p, "{% cellbg task.bg %}", size=8)
        add_run(p, tag, size=8)

    # {%tr endfor %}
    p = table.cell(3, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr endfor %}", size=8)
    for i in range(1, 3):
        set_paragraph_spacing(table.cell(3, i).paragraphs[0], before=0, after=0)


def build_metadata_grid(doc):
    """Build the 2x3 metadata grid table."""
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color=FINE_BORDER, sz=2)

    grid = [
        [("CASE NO.", "{{ case.case_number }}", "Consolas"),
         ("JUDGE", "{{ case.judge }}", "Calibri"),
         ("DATE OF INJURY", "{{ case.doi }}", "Calibri")],
        [("JURISDICTION", "{{ case.jurisdiction }}", "Calibri"),
         ("OPP. COUNSEL", "{{ case.opp_counsel }}", "Calibri"),
         ("CLIENTS", "{{ case.clients }}", "Calibri")],
    ]

    col_width = 2.4

    for row_idx, row_data in enumerate(grid):
        for col_idx, (label, tag, font) in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            set_cell_shading(cell, LIGHT_BG)
            set_cell_margins(cell, top=20, bottom=20, left=55, right=55)
            set_cell_width(cell, col_width)

            p_label = cell.paragraphs[0]
            set_paragraph_spacing(p_label, before=0, after=0)
            add_run(p_label, label, size=6, bold=True, color=GREY)

            p_val = cell.add_paragraph()
            set_paragraph_spacing(p_val, before=0, after=0)
            add_run(p_val, tag, size=8.5, font_name=font)


def build_title_bar(doc):
    """Build title bar: case name left, status chip right."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    remove_table_borders(table)

    # Col 0: case name
    left = table.cell(0, 0)
    set_cell_width(left, 5.2)
    set_cell_margins(left, top=0, bottom=8, left=0, right=0)
    p = left.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{{r case.title_rt }}", size=14)

    # Col 1: spacer
    spacer = table.cell(0, 1)
    set_cell_width(spacer, 0.4)
    set_cell_margins(spacer, top=0, bottom=0, left=0, right=0)
    set_paragraph_spacing(spacer.paragraphs[0], before=0, after=0)

    # Col 2: status chip — compact colored cell
    chip = table.cell(0, 2)
    set_cell_width(chip, 1.6)
    set_cell_margins(chip, top=40, bottom=10, left=50, right=50)
    # All-sides border in the status bg color — creates a filled chip look
    set_cell_borders(chip,
                     top=("E2E8F0", 4),
                     bottom=("E2E8F0", 4),
                     left=("E2E8F0", 4),
                     right=("E2E8F0", 4))
    p = chip.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{% cellbg case.status_bg %}", size=1, color=WHITE)
    add_run(p, "{{r case.status_rt }}", size=10)

    # Accent rule below title bar
    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=0, after=0)
    set_paragraph_border_bottom(rule, color=ACCENT_HEX, sz=6)


def build_case_summary(doc):
    """Build case summary as a paragraph with blue left border."""
    # Conditional: only show if summary exists
    if_p = doc.add_paragraph()
    set_paragraph_spacing(if_p, before=0, after=0)
    add_run(if_p, "{% if case.summary %}", size=1, color=WHITE)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=3)
    # Blue left border
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBorders")
    left_b = OxmlElement("w:left")
    left_b.set(qn("w:val"), "single")
    left_b.set(qn("w:sz"), "14")
    left_b.set(qn("w:color"), ACCENT_HEX)
    left_b.set(qn("w:space"), "6")
    borders.append(left_b)
    p_pr.append(borders)
    # Indent the paragraph slightly
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "120")
    p_pr.append(ind)
    add_run(p, "{{ case.summary }}", size=8, italic=True, color=DARK_GREY)

    end_p = doc.add_paragraph()
    set_paragraph_spacing(end_p, before=0, after=0)
    add_run(end_p, "{% endif %}", size=1, color=WHITE)


def build_timeline_table(doc):
    """Build side-by-side events + tasks table with visual divider."""
    # 4 data columns + 1 divider column = 5 cols total
    table = doc.add_table(rows=5, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color=FINE_BORDER, sz=2)

    # Col widths: ev_date, ev_desc, DIVIDER, tk_due, tk_desc
    col_widths = [1.2, 2.35, 0.1, 1.0, 2.55]

    # Row 0: Section headers — events (cols 0-1), divider (col 2), tasks (cols 3-4)
    ev_header = table.cell(0, 0).merge(table.cell(0, 1))
    set_cell_shading(ev_header, HEADER_BG)
    set_cell_margins(ev_header, top=22, bottom=22, left=50, right=50)
    p = ev_header.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "KEY DATES & EVENTS", size=7, bold=True, color=WHITE)

    div_header = table.cell(0, 2)
    set_cell_shading(div_header, "FFFFFF")
    set_cell_width(div_header, 0.1)
    set_cell_margins(div_header, top=0, bottom=0, left=0, right=0)
    set_paragraph_spacing(div_header.paragraphs[0], before=0, after=0)

    tk_header = table.cell(0, 3).merge(table.cell(0, 4))
    set_cell_shading(tk_header, HEADER_BG)
    set_cell_margins(tk_header, top=22, bottom=22, left=50, right=50)
    p = tk_header.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "OPEN TASKS", size=7, bold=True, color=WHITE)

    # Row 1: Sub-headers
    sub_headers = ["Date", "Description", "", "Due", "Description"]
    for i, h in enumerate(sub_headers):
        cell = table.cell(1, i)
        if i == 2:  # divider
            set_cell_shading(cell, "FFFFFF")
            set_cell_width(cell, 0.1)
            set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
        else:
            set_cell_shading(cell, LIGHT_BG)
            set_cell_margins(cell, top=16, bottom=16, left=50, right=50)
            set_cell_width(cell, col_widths[i])
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        if h:
            add_run(p, h, size=6, bold=True, color=GREY)

    # Row 2: {%tr for item in case.timeline %}
    p = table.cell(2, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr for item in case.timeline %}", size=8)
    for i in range(1, 5):
        set_paragraph_spacing(table.cell(2, i).paragraphs[0], before=0, after=0)

    # Row 3: Data row
    data_tags = [
        ("{{r item.ev_date }}", col_widths[0]),
        ("{{r item.ev_desc }}", col_widths[1]),
        ("", col_widths[2]),  # divider
        ("{{r item.tk_date }}", col_widths[3]),
        ("{{r item.tk_desc }}", col_widths[4]),
    ]
    for i, (tag, w) in enumerate(data_tags):
        cell = table.cell(3, i)
        set_cell_width(cell, w)
        if i == 2:  # divider column
            set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
            set_cell_shading(cell, "FFFFFF")
            set_paragraph_spacing(cell.paragraphs[0], before=0, after=0)
        else:
            set_cell_margins(cell, top=12, bottom=12, left=50, right=50)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0)
            add_run(p, "{% cellbg item.bg %}", size=1, color=WHITE)
            if tag:
                add_run(p, tag, size=8)

    # Row 4: {%tr endfor %}
    p = table.cell(4, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0)
    add_run(p, "{%tr endfor %}", size=8)
    for i in range(1, 5):
        set_paragraph_spacing(table.cell(4, i).paragraphs[0], before=0, after=0)


def build_case_detail(doc):
    """Build per-case detail: title+chip, metadata, summary, timeline, notes."""
    # --- Outer loop: phases ---
    outer_start = doc.add_paragraph()
    set_paragraph_spacing(outer_start, before=0, after=0)
    add_run(outer_start, "{% for phase in phases %}", size=1, color=WHITE)

    # --- Inner loop: cases in phase ---
    inner_start = doc.add_paragraph()
    set_paragraph_spacing(inner_start, before=0, after=0)
    add_run(inner_start, "{% for case in phase.cases %}", size=1, color=WHITE)

    # Page break
    pb_p = doc.add_paragraph()
    set_paragraph_spacing(pb_p, before=0, after=0)
    set_page_break_before(pb_p)
    add_run(pb_p, " ", size=1, color=WHITE)

    # --- 1. Title bar (case name + status chip) ---
    build_title_bar(doc)

    # --- 2. Metadata grid (2x3) — tight to title ---
    build_metadata_grid(doc)

    # --- 3. Summary (below metadata, conditional) ---
    build_case_summary(doc)

    # --- 4. Other Proceedings (conditional, compact) ---
    op_p = doc.add_paragraph()
    set_paragraph_spacing(op_p, before=0, after=0)
    add_run(op_p, "{% if case.other_proceedings %}", size=1, color=WHITE)

    op_content = doc.add_paragraph()
    set_paragraph_spacing(op_content, before=1, after=1)
    add_run(op_content, "Other Proceedings: ", size=7.5, bold=True, color=NAVY)
    add_run(op_content, "{{ case.other_proceedings }}", size=7.5, color=DARK_GREY, font_name="Consolas")

    op_end = doc.add_paragraph()
    set_paragraph_spacing(op_end, before=0, after=0)
    add_run(op_end, "{% endif %}", size=1, color=WHITE)

    # --- 5. Timeline (events + tasks side by side) ---
    tl_if = doc.add_paragraph()
    set_paragraph_spacing(tl_if, before=2, after=0)
    add_run(tl_if, "{% if case.has_timeline %}", size=1, color=WHITE)

    build_timeline_table(doc)

    tl_end = doc.add_paragraph()
    set_paragraph_spacing(tl_end, before=0, after=0)
    add_run(tl_end, "{% endif %}", size=1, color=WHITE)

    # --- 6. Recent Notes ---
    nt_if = doc.add_paragraph()
    set_paragraph_spacing(nt_if, before=0, after=0)
    add_run(nt_if, "{% if case.notes %}", size=1, color=WHITE)

    build_section_heading(doc, "RECENT NOTES")

    nt_loop = doc.add_paragraph()
    set_paragraph_spacing(nt_loop, before=0, after=0)
    add_run(nt_loop, "{% for note in case.notes %}", size=1, color=WHITE)

    note_date_p = doc.add_paragraph()
    set_paragraph_spacing(note_date_p, before=2, after=0)
    add_run(note_date_p, "{{ note.date }}", size=7, bold=True, color=GREY)

    note_content_p = doc.add_paragraph()
    set_paragraph_spacing(note_content_p, before=0, after=2)
    p_pr2 = note_content_p._p.get_or_add_pPr()
    borders2 = OxmlElement("w:pBorders")
    left_b2 = OxmlElement("w:left")
    left_b2.set(qn("w:val"), "single")
    left_b2.set(qn("w:sz"), "10")
    left_b2.set(qn("w:color"), ACCENT_HEX)
    left_b2.set(qn("w:space"), "6")
    borders2.append(left_b2)
    p_pr2.append(borders2)
    add_run(note_content_p, "{{ note.content }}", size=8, color=DARK_GREY)

    nt_loop_end = doc.add_paragraph()
    set_paragraph_spacing(nt_loop_end, before=0, after=0)
    add_run(nt_loop_end, "{% endfor %}", size=1, color=WHITE)

    nt_end = doc.add_paragraph()
    set_paragraph_spacing(nt_end, before=0, after=0)
    add_run(nt_end, "{% endif %}", size=1, color=WHITE)

    # --- End inner case loop ---
    inner_end = doc.add_paragraph()
    set_paragraph_spacing(inner_end, before=0, after=0)
    add_run(inner_end, "{% endfor %}", size=1, color=WHITE)

    # --- End outer phase loop ---
    outer_end = doc.add_paragraph()
    set_paragraph_spacing(outer_end, before=0, after=0)
    add_run(outer_end, "{% endfor %}", size=1, color=WHITE)


def build():
    """Build the complete template document."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(9)
    font.color.rgb = NAVY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Header & footer
    build_header_footer(doc)

    # Cover page
    build_cover_page(doc)

    # Per-case detail pages
    build_case_detail(doc)

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Template saved to {OUTPUT}")


if __name__ == "__main__":
    build()
