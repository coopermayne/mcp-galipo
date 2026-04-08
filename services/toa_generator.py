"""
Table of Authorities document generator.

Takes extracted authorities and generates a properly formatted TOA as a .docx file
using python-docx. Port of the spec's JavaScript generation logic.
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER


FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
LINE_SPACING = Pt(24)  # Double-spaced


def _format_pages(pages: list[int], offset: int = 0) -> str:
    """
    Format page numbers with offset, collapsing consecutive runs into ranges.

    6+ unique pages becomes "passim".
    3+ consecutive pages become an en-dash range.
    """
    if not pages:
        return "passim"

    adjusted = sorted(set(p + offset for p in pages))

    if len(adjusted) >= 6:
        return "passim"

    # Group consecutive pages into ranges
    ranges = []
    start = end = adjusted[0]
    for p in adjusted[1:]:
        if p == end + 1:
            end = p
        else:
            ranges.append((start, end))
            start = end = p
    ranges.append((start, end))

    parts = []
    for s, e in ranges:
        if s == e:
            parts.append(str(s))
        elif e - s >= 2:
            parts.append(f"{s}\u2013{e}")  # en-dash for 3+ consecutive
        else:
            parts.append(f"{s}, {e}")

    return ", ".join(parts)


def _add_run(paragraph, text, bold=False, italic=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    run.italic = italic
    return run


def _set_paragraph_format(paragraph, spacing=LINE_SPACING, alignment=None,
                          indent_left=None, indent_hanging=None,
                          space_before=None, space_after=None):
    """Configure paragraph formatting."""
    fmt = paragraph.paragraph_format
    fmt.line_spacing = spacing
    if alignment is not None:
        fmt.alignment = alignment
    if indent_left is not None:
        fmt.left_indent = indent_left
    if indent_hanging is not None:
        fmt.first_line_indent = -indent_hanging
    if space_before is not None:
        fmt.space_before = space_before
    if space_after is not None:
        fmt.space_after = space_after


def _add_dot_leader_tab(paragraph):
    """Add a right-aligned dot-leader tab stop at the right margin."""
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(
        Inches(6.5),  # 6.5 inches (page width 8.5 - 1 inch margins)
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.DOTS,
    )


# Category display order and labels
CATEGORY_ORDER = [
    ("case", "Cases"),
    ("constitutional_provision", "Constitutional Provisions"),
    ("statute", "Statutes"),
    ("rule", "Rules"),
    ("other", "Other Authorities"),
]


def generate_toa(authorities: list[dict], page_offset: int = 0) -> BytesIO:
    """
    Generate a Table of Authorities .docx document.

    Args:
        authorities: List of authority dicts from extraction
        page_offset: Integer to add to all page numbers

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # Set page margins (1 inch all around)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title_para = doc.add_paragraph()
    _set_paragraph_format(
        title_para,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12),
    )
    _add_run(title_para, "TABLE OF AUTHORITIES", bold=True)

    # Group by category
    for cat_key, cat_label in CATEGORY_ORDER:
        entries = [a for a in authorities if a.get("category") == cat_key]
        if not entries:
            continue

        # Sort entries
        if cat_key == "case":
            entries.sort(key=lambda a: (a.get("name") or "").lower())
        elif cat_key == "other":
            entries.sort(key=lambda a: (a.get("author") or "").lower())
        else:
            entries.sort(key=lambda a: (a.get("text") or "").lower())

        # Category heading
        heading_para = doc.add_paragraph()
        _set_paragraph_format(
            heading_para,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(18),
            space_after=Pt(12),
        )
        run = _add_run(heading_para, cat_label, bold=True)
        run.underline = True

        # Entries
        for entry in entries:
            page_str = _format_pages(entry.get("pages", []), page_offset)

            para = doc.add_paragraph()
            _set_paragraph_format(
                para,
                indent_left=Inches(0.5),
                indent_hanging=Inches(0.5),
            )
            _add_dot_leader_tab(para)

            if cat_key == "case":
                # Case: italic name, comma, line break, roman cite with tab + pages
                _add_run(para, entry.get("name", ""), italic=True)
                _add_run(para, ",")
                # Add a line break then the cite
                run = para.add_run()
                run.add_break()
                _add_run(para, f"{entry.get('cite', '')}\t")
                _add_run(para, page_str)

            elif cat_key == "other":
                # Other: author (roman), title (italic), rest (roman) + tab + pages
                _add_run(para, entry.get("author", ""))
                _add_run(para, entry.get("title", ""), italic=True)
                _add_run(para, f"{entry.get('rest', '')}\t")
                _add_run(para, page_str)

            else:
                # Statutes, constitutional provisions, rules: all roman
                _add_run(para, f"{entry.get('text', '')}\t")
                _add_run(para, page_str)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
